import datetime
from django.forms.models import model_to_dict
import bangla
from django.contrib import messages

import os

from bs4 import BeautifulSoup
from PIL import Image
from PIL import ImageDraw
from PIL import ImageFont
from docx import Document

from ac_openning.models import Customer, Account, TxnProfile, Occupation, SourceOfIncome
from ac_opn_nid.settings import NID_DIRECTORY, N_NID_DIRECTORY, STATIC_URL
from helper.helper import CONVERSION_VAR


def process_basic_info(item):
    basic_info = {"Name(Bangla)": "b_name",
                  "Name(English)": "e_name",
                  "Date of Birth": "dob",
                  "Father Name": "f_name",
                  "Mother Name": "m_name",
                  "Spouse Name": "s_name",
                  "Occupation": "occupation_nid"}
    processed_info = {}

    for itm in basic_info.items():
        val_div = item[0].find_all('div', string=itm[0])
        val = val_div[0].find_next_sibling()

        processed_info[itm[1]] = val.text
    return processed_info


def process_pre_address(pre_div):
    pre_info = {"Division": "pre_division",
                "District": "pre_district",
                "RMO": "pre_rmo",
                "City Corporation Or Municipality": "pre_citycorp",
                "Upozila": "pre_upozila",
                "Union/Ward": "pre_union_ward",
                "mouza/Moholla": "pre_mouza_moholla",
                "Additional Mouza/Moholla": "pre_additional_mouza_moholla",
                "Ward For Union Porishod": "pre_ward_for_union_porishod",
                "village/Road": "pre_village_road",
                "Additional Village/Road": "pre_additional_village_road",
                "Home/Holding No": "pre_home_holding_no",
                "Post Office": "pre_post_office",
                "Postal Code": "pre_postal_code",
                "Region": "pre_region"}
    processed_info = {}

    # import pdb;pdb.set_trace()
    for itm in pre_info.items():
        val_div = pre_div.find_all('div', string=itm[0])
        val = val_div[0].find_next_sibling()

        processed_info[itm[1]] = val.text
    return processed_info


def process_per_address(per_div):
    per_info = {"Division": "per_division",
                "District": "per_district",
                "RMO": "per_rmo",
                "City Corporation Or Municipality": "per_city_corporation_or_municipality",
                "Upozila": "per_upozila",
                "Union/Ward": "per_union_ward",
                "mouza/Moholla": "per_mouza_moholla",
                "Additional Mouza/Moholla": "per_additional_mouza_moholla",
                "Ward For Union Porishod": "per_ward_for_union_porishod",
                "village/Road": "per_village_road",
                "Additional Village/Road": "per_additional_village_road",
                "Home/Holding No": "per_home_holding_no",
                "Post Office": "per_post_office",
                "Postal Code": "per_postal_code",
                "Region": "per_region"}
    processed_info = {}

    for itm in per_info.items():
        val_div = per_div.find_all('div', string=itm[0])
        val = val_div[0].find_next_sibling()

        processed_info[itm[1]] = val.text
    return processed_info


def process_other_info(item):
    basic_info = {"Blood Group": "blood_group",
                  "National ID": "national_id",
                  "Pin": "pin"}
    processed_info = {}

    for itm in basic_info.items():
        val_div = item[0].find_all('div', string=itm[0])
        val = val_div[0].find_next_sibling()

        processed_info[itm[1]] = val.text
    return processed_info


def process_data(file_name):
    with open(NID_DIRECTORY + file_name, encoding="utf8") as fp:
        soup = BeautifulSoup(fp, 'html.parser')
        item = soup.find_all('div', {'class': 'row border'})
        processed_info = process_basic_info(item)

        per_div = item[0].find_all('div', string='Permanent Address')
        per_div = per_div[0].find_next_sibling()
        processed_info = {**processed_info, **process_per_address(per_div)}

        pre_div = item[0].find_all('div', string='Present Address')
        pre_div = pre_div[0].find_next_sibling()
        processed_info = {**processed_info, **process_pre_address(pre_div)}
        processed_info = {**processed_info, **process_other_info(item)}
        return processed_info


def update_or_insert_customer(data):
    customer = Customer.objects.filter(national_id=data.get("national_id"))
    if customer.exists():
        customer = Customer.objects.get(national_id=data.get("national_id"))
        for item in data.items():
            if item[0] is not "national_id":
                setattr(customer, item[0], item[1])
        message = "Updated customer successfully"

    else:
        customer = Customer(**data)
        message = "Created customer successfully"
    customer.save()

    return message


def handle_uploaded_file(fl, name):
    dir = NID_DIRECTORY

    if os.path.exists(dir + name):
        os.remove(dir + name)

    try:
        with open(NID_DIRECTORY + name, "wb+") as destination:
            for chunk in fl.chunks():
                destination.write(chunk)
            destination.close()
    except Exception as ex:
        print(str(ex))
        # import pdb;pdb.set_trace()


def fill_up_form(acc_id, cust_id, tp_id, nom_id=None):
    res ={}
    doc = Document('test.docx')

    account = model_to_dict(Account.objects.get(id=acc_id))
    customer = model_to_dict(Customer.objects.get(id=cust_id))

    tp = model_to_dict(TxnProfile.objects.get(id=tp_id))
    if nom_id is not None:
        nominee = model_to_dict(Customer.objects.get(id=nom_id))
        mod_nominee = {}
        for k,v in nominee.items():
            mod_nominee["n_"+k] = nominee[k]
        res = mod_nominee
    account["ac_b_name"] = account["b_name"]
    account["ac_e_name"] = account["e_name"]

    account.pop("b_name")
    account.pop("e_name")

    res = { **res, **account, **customer, **tp}

    # import pdb;pdb.set_trace()

    # for paragraph in doc.paragraphs:
    #     for doc_key, dict_key in CONVERSION_VAR.items():
    #         if doc_key in doc.paragraphs:
    #             paragraph.text = paragraph.text.replace(doc_key, res.get(dict_key))
    dtot_cnt = 0
    dtot_amt = 0
    for tab in doc.tables:
        for row in tab.rows:
            for cell in row.cells:
                if "__dccnt__" in cell.text:
                    cell.text = cell.text.replace("__dccnt__", res.get("depo_cash_txn"))
                if "__dtcnt__" in cell.text:
                    cell.text = cell.text.replace("__dtcnt__", res.get("depo_trans_txn"))
                if "__dcamt__" in cell.text:
                    cell.text = cell.text.replace("__dcamt__", res.get("depo_cash_txn_amt"))
                if "__dtamt__" in cell.text:
                    cell.text = cell.text.replace("__dtamt__", res.get("depo_trans_txn_amt"))
                if "__wdccnt__" in cell.text:
                    cell.text = cell.text.replace("__wdccnt__", res.get("wd_cash_txn"))
                if "__wdtcnt__" in cell.text:
                    cell.text = cell.text.replace("__wdtcnt__", res.get("wd_trans_txn"))
                if "__wdcamt__" in cell.text:
                    cell.text = cell.text.replace("__wdcamt__", res.get("wd_cash_txn_amt"))
                if "__wdtamt__" in cell.text:
                    cell.text = cell.text.replace("__wdtamt__", res.get("wd_trans_txn_amt"))
    # import pdb;pdb.set_trace()
    for paragraph in doc.paragraphs:
        if '__b_name__' in paragraph.text:
            paragraph.text = paragraph.text.replace("__b_name__", res.get('b_name', " "))
        if "__ac_b_name__" in paragraph.text:
            paragraph.text = paragraph.text.replace("__ac_b_name__", res.get("ac_b_name", " "))
        if "__ac_e_name__" in paragraph.text:
            paragraph.text = paragraph.text.replace("__ac_e_name__", res.get("ac_e_name", " "))
        if "__init_depo__" in paragraph.text:
            paragraph.text = paragraph.text.replace("__init_depo__", res.get("init_depo", " "))
        if "__init_depo_word__" in paragraph.text:
            paragraph.text = paragraph.text.replace("__init_depo_word__", res.get("init_depo_word", " "))
        if "__b_name__" in paragraph.text:
            paragraph.text = paragraph.text.replace("__b_name__", res.get("b_name", " "))
        if "__e_name__" in paragraph.text:
            paragraph.text = paragraph.text.replace("__e_name__", res.get("e_name", " "))
        if "__dob__" in paragraph.text:
            paragraph.text = paragraph.text.replace("__dob__", res.get("dob", " "))
        if "__f_name__" in paragraph.text:
            paragraph.text = paragraph.text.replace(" __f_name__", res.get("f_name", " "))
        if "__m_name__" in paragraph.text:
            paragraph.text = paragraph.text.replace("__m_name__", res.get("m_name", " "))
        if "__s_name__" in paragraph.text:
            paragraph.text = paragraph.text.replace(" __s_name__", res.get("s_name", " "))
        if "__nationality__" in paragraph.text:
            paragraph.text = paragraph.text.replace("__nationality__", res.get("nationality", " "))
        if "__vill__" in paragraph.text:
            val = res.get('pre_village_road', 'pre_additional_village_road')
            paragraph.text = paragraph.text.replace("__vill__", val) if val else " "
        if "__post__" in paragraph.text:
            paragraph.text = paragraph.text.replace("__post__", res.get("pre_post_office", " "))
        if "__thana__" in paragraph.text:
            paragraph.text = paragraph.text.replace("__thana__", res.get("pre_upozila", " "))
        if "__dist__" in paragraph.text:
            paragraph.text = paragraph.text.replace("__dist__", res.get("pre_district", " "))
        if "__pvill__" in paragraph.text:
            val = res.get('per_village_road', 'per_additional_village_road')
            paragraph.text = paragraph.text.replace("__pvill__", val) if val else " "
        if "__ppost__" in paragraph.text:
            paragraph.text = paragraph.text.replace("__ppost__", res.get("per_post_office", " "))
        if "__pthana__" in paragraph.text:
            paragraph.text = paragraph.text.replace("__pthana__", res.get("per_upozila", " "))
        if "__pdist__" in paragraph.text:
            paragraph.text = paragraph.text.replace("__pdist__", res.get("per_district", " "))
        if "__n_name__" in paragraph.text:
            paragraph.text = paragraph.text.replace("__n_name__", res.get("n_b_name", " "))
        if "__n_dob__" in paragraph.text:
            paragraph.text = paragraph.text.replace("__n_dob__", res.get("n_dob", " "))
        if "__n_f_name__" in paragraph.text:
            paragraph.text = paragraph.text.replace("__n_f_name__", res.get("n_f_name", " "))
        if "__n_m_name__" in paragraph.text:
            paragraph.text = paragraph.text.replace("__n_m_name__", res.get("n_m_name", " "))
        if "__n_vill__" in paragraph.text:
            val = res.get('n_pre_village_road', 'n_pre_additional_village_road')
            paragraph.text = paragraph.text.replace("__n_vill__", val) if val else " "
        if "__n_post__" in paragraph.text:
            paragraph.text = paragraph.text.replace("__n_post__", res.get("n_pre_post_office", " "))
        if "__n_thana__" in paragraph.text:
            paragraph.text = paragraph.text.replace("__n_thana__", res.get("n_pre_upozila", " "))
        if "__n_dist__" in paragraph.text:
            paragraph.text = paragraph.text.replace("__n_dist__", res.get("n_pre_district", " "))
        if "__n_national_id__" in paragraph.text:
            paragraph.text = paragraph.text.replace("__n_national_id__", res.get("n_national_id", " "))
        if "__occupation__" in paragraph.text:
            paragraph.text = paragraph.text.replace("__occupation__", res.get("occupation", " "))
        if "__income__" in paragraph.text:
            paragraph.text = paragraph.text.replace("__income__", res.get("income", " "))
        if "__source__" in paragraph.text:
            paragraph.text = paragraph.text.replace("__source__", res.get("source", " "))
        if "__national_id__" in paragraph.text:
            paragraph.text = paragraph.text.replace("__national_id__", res.get("national_id", " "))
        if "__dtot_cnt__" in paragraph.text:
            paragraph.text = paragraph.text.replace("__dtot_cnt__", " ")
        if "__dtot_amt__" in paragraph.text:
            paragraph.text = paragraph.text.replace("__dtot_amt__", " ")
        if "__dtot_cnt__" in paragraph.text:
            paragraph.text = paragraph.text.replace("__dtot_cnt__", " ")
        if "__dtot_amt__" in paragraph.text:
            paragraph.text = paragraph.text.replace("__dtot_amt__", " ")
        if "__date__" in paragraph.text:
            paragraph.text = paragraph.text.replace("__date__", str(datetime.datetime.today().date()))

    # filepath = STATIC_URL+ "download/" +res.get("national_id") + ".docx"
    filepath = res.get("national_id") +".docx"
    doc.save(filepath)

    return filepath


def process_for_show_data(nid, nnid=None):
    customer = Customer.objects.get(national_id=nid)
    source =SourceOfIncome.objects.all()
    occupation =Occupation.objects.all()
    cust_dict = model_to_dict(customer)
    nom_dict = {}
    acc_dict = {}
    tp_dict = {}
    if nnid is not None:
        nominee = Customer.objects.get(national_id=nnid)
        nom_dict = model_to_dict(nominee)
    fields = Account._meta.get_fields()
    for item in fields:
        if item.name != 'customer' and item.name != 'nominee' and item.name != 'txnprofile':
            acc_dict[item.name] = ""
    fields = TxnProfile._meta.get_fields()
    for item in fields:
        if item.name != 'account':
            tp_dict[item.name] = ""

    # import pdb;pdb.set_trace()
    return {"customer": cust_dict, "nominee": nom_dict, "acc_info": acc_dict, "tp": tp_dict, "source": source, "occupation": occupation}


def process_for_form_fillup(data):
    acc_dict={}
    cust_dict={}
    nom_dict={}
    tp_dict = {}
    for k,v in data.items():
        if 'acc_' in k[0:4]:
            acc_dict[k.replace('acc_', '')] = v
        elif 'c_' in k[0:4]:
            cust_dict[k.replace('c_', '')] = v
        elif 'n_' in k[0:4]:
            nom_dict[k.replace('n_', '')] = v
        else:
            tp_dict[k.replace('tp_', '')] = v
    return acc_dict, cust_dict, nom_dict, tp_dict


def update_db(acc_dict, cust_dict, nom_dict, tp_dict):
    customer = Customer.objects.get(national_id=cust_dict.get("national_id"))
    nom_id = None
    acc = Account.objects.filter(ac_number=acc_dict.get("ac_number"))
    if acc.exists():
        acc = Account.objects.get(ac_number=acc_dict.get("ac_number"))
        acc_dict.pop("id")

        if bool(nom_dict) is not False:
            # import pdb;pdb.set_trace()
            nominee = Customer.objects.get(national_id=nom_dict.get("national_id"))
            nom_id = nominee.id
            acc_dict["nominee"] = nom_id
        for item in acc_dict.items():
            setattr(acc, item[0], item[1])
        message = "Updated Acc successfully"

    else:

        acc_dict.pop("id")
        acc_dict["customer_id"] = customer.id
        acc = Account(**acc_dict)
        message = "Created Acc successfully"
    acc.save()

    tp = TxnProfile.objects.filter(account_id=acc.id)
    if tp.exists():
        tp = TxnProfile.objects.get(account_id=acc.id)
        tp_dict.pop("id")
        for item in tp_dict.items():
            setattr(tp, item[0], item[1])
        message = "Updated tp successfully"

    else:
        tp_dict.pop("id")
        tp_dict["account_id"] = acc.id
        tp = TxnProfile(**tp_dict)
        message = "Created tp successfully"
    tp.save()

    return acc.id, customer.id, nom_id, tp.id
